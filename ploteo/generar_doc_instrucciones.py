#!/usr/bin/env python3
"""Genera INSTRUCCIONES_INSTALACION.docx con la guía de instalación de NewPT."""

import os
import sys
from datetime import date

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:
    print("Falta el módulo 'docx'. Instálalo y vuelve a ejecutar:")
    print("    pip install python-docx")
    sys.exit(1)

DIR = os.path.dirname(os.path.abspath(__file__))
SALIDA = os.path.join(DIR, "INSTRUCCIONES_INSTALACION.docx")


def codigo(doc, texto):
    """Bloque de comando en fuente monoespaciada."""
    for linea in texto.rstrip("\n").split("\n"):
        p = doc.add_paragraph()
        r = p.add_run(linea)
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)


def nota(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run("NOTA: " + texto)
    r.bold = True
    r.font.size = Pt(10)
    return p


def main():
    doc = Document()
    doc.add_heading("Instrucciones de instalación de NewPT", 0)
    doc.add_paragraph(
        "Guía completa para instalar el proyecto NewPT (contexto sismotectónico de eventos "
        "sísmicos) en una máquina nueva, sin permisos de super usuario y sin instalar módulos "
        "de Python en el Python del sistema."
    )
    doc.add_paragraph("Generado el " + date.today().strftime("%d/%m/%Y") + ".")

    # ------------------------------------------------------------------
    doc.add_heading("1. Qué hace el instalador", 1)
    doc.add_paragraph(
        "El instalador (instalar_newpt.sh) compila los scripts a binarios con PyInstaller "
        "dentro de un ambiente virtual (venv) y deja una instalación autocontenida. Todo se "
        "hace como usuario normal, sin sudo:"
    )
    for t in [
        "No instala ningún módulo de Python en el Python del sistema (todo va en el venv).",
        "Los códigos fuente y el venv quedan ocultos en la carpeta .dev/ (chmod 700).",
        "Genera binarios ejecutables de consulta y de graficado.",
        "Copia los datos (grillas, TIF, sismicidad histórica, localidades).",
        "Crea el lanzador newpt.sh y la configuración de base de datos (.env).",
        "El instalador NUNCA ejecuta apt: los programas del sistema solo se verifican, "
        "mostrando por pantalla el estado de cada uno ([OK] o [DEBE INSTALARSE]). Si falta "
        "alguno, muestra un resumen con los pendientes y DETIENE la instalación.",
        "Los módulos de Python se instalan en el venv con versiones EXACTAMENTE FIJADAS "
        "(las mismas probadas en el desarrollo) y al final imprime las versiones instaladas.",
        "Antes de compilar los binarios ejecuta la MARCHA BLANCA: pruebas unitarias "
        "(pytest) sobre la lógica crítica y los datos. Si un test falla, la instalación "
        "se DETIENE. Detalle completo en INFORME_TESTS_MARCHA_BLANCA.docx.",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    # ------------------------------------------------------------------
    doc.add_heading("2. Rutina paso a paso para cada máquina nueva", 1)

    doc.add_heading("Paso 1 - En la máquina de desarrollo (carpeta ploteo/)", 2)
    doc.add_paragraph("Solo si cambiaste código, verifica que no haya errores de sintaxis:")
    codigo(doc, "python3 -m py_compile capturar.py consulta_evento.py consuta_evento_sc7.py")
    doc.add_paragraph("Regenera el paquete de instalación (fuentes + datos + herramientas):")
    codigo(doc, "./crear_instalador.sh")
    doc.add_paragraph("El resultado es newpt_instalador.tar.gz en la carpeta actual "
                      "(o en la que indiques como argumento).")
    doc.add_paragraph("Copia el paquete a la máquina nueva (scp, USB, sftp, rsync, etc.):")
    codigo(doc, "scp newpt_instalador.tar.gz usuario@maquina:~/")

    doc.add_heading("Paso 2 - En la máquina nueva (usuario normal)", 2)
    codigo(
        doc,
        "tar -xzf newpt_instalador.tar.gz\n"
        "cd newpt_instalador\n"
        "./instalar_newpt.sh --solo-verificar   # opcional: comprueba prerequisitos\n"
        "./instalar_newpt.sh --probar           # opcional: marcha blanca (tests)\n"
        "./instalar_newpt.sh                    # instala en ~/newpt\n"
        "~/newpt/newpt.sh csn_sc62026nkkbb      # prueba con el ID de un evento real",
    )
    doc.add_paragraph(
        "El modo --probar crea un venv temporal, instala los módulos con versiones fijas, "
        "corre las pruebas unitarias (ver INFORME_TESTS_MARCHA_BLANCA.docx) y borra lo "
        "temporal: no instala nada en ~/newpt."
    )
    doc.add_paragraph(
        "En la prueba, reemplace 'csn_sc62026nkkbb' por el ID del evento que quiera graficar "
        "y escríbalo tal cual (sin signos < > ni comillas): los signos angulares de otros "
        "textos solo indican dónde va un valor, no se escriben en la terminal."
    )
    doc.add_paragraph(
        "La instalación compila los binarios, por lo que puede tardar varios minutos "
        "y necesita acceso a internet (o a un mirror de pip) para descargar los módulos."
    )

    doc.add_heading("Paso 3 - Prerequisitos del sistema (SOLO la primera vez por máquina)", 2)
    doc.add_paragraph("Estos programas se solicitan al encargado de TI y requieren sudo:")
    codigo(doc, "sudo apt update\nsudo apt install -y python3 python3-venv python3-tk xterm")
    doc.add_paragraph(
        "En una máquina donde NewPT ya estuvo instalado o corriendo antes, estos programas "
        "ya existen: NO se vuelven a instalar. El instalador verifica cada uno y muestra su "
        "estado ([OK] o [DEBE INSTALARSE]). Si alguno falta, lista los pendientes y DETIENE "
        "la instalación sin continuar; una vez instalados por TI, basta volver a ejecutarlo."
    )

    doc.add_heading("Paso 4 - Configurar el botón de SeisComP", 2)
    doc.add_paragraph("El botón debe ejecutar el lanzador con el ID del evento como argumento:")
    codigo(doc, "~/newpt/newpt.sh $EVENT_ID")
    doc.add_paragraph(
        "OJO: en este caso SÍ se escribe literalmente $EVENT_ID, porque SeisComP lo reemplaza "
        "automáticamente por el ID del evento al presionar el botón. No confundir con la "
        "prueba manual del Paso 2, donde se escribe el ID real del evento."
    )
    doc.add_paragraph("Ajusta la ruta según la carpeta elegida en la instalación.")

    # ------------------------------------------------------------------
    doc.add_heading("3. Máquinas que ya corrían el proyecto con los códigos originales", 1)
    for t in [
        "No hay conflicto: el instalador crea una instalación independiente en ~/newpt.",
        "Los códigos y datos originales NO se tocan, quedan donde estaban.",
        "Los programas del sistema (python3, python3-tk, xterm) ya están instalados: "
        "no se vuelven a instalar.",
        "Si instalas en una carpeta que ya contenía una instalación previa de NewPT "
        "(newpt.sh / bin/ / .dev/), el instalador lo detecta y reconstruye .dev/ y bin/ "
        "de forma limpia, conservando el resto.",
        "Al terminar, la carpeta antigua puede ignorarse o borrarse a mano, y el botón de "
        "SeisComP se apunta al nuevo lanzador.",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    # ------------------------------------------------------------------
    doc.add_heading("4. Estructura de la instalación (~/newpt)", 1)
    codigo(
        doc,
        "~/newpt/\n"
        "├── newpt.sh                     lanzador (lo ejecuta el botón de SeisComP)\n"
        "├── .env                         configuración de la base de datos (chmod 600)\n"
        "├── bin/newpt_capturar/          binario del graficador\n"
        "├── bin/newpt_consulta/          binario de la consulta\n"
        "├── grillas/                     perfiles de subducción y topografía\n"
        "├── NE2_LR_LC_SR_W_DR.tif        imagen de relieve\n"
        "├── base_2023_2026.dat           sismicidad histórica\n"
        "├── localidades.csv              localidades\n"
        "└── .dev/                        venv + fuentes (oculto, chmod 700)",
    )

    # ------------------------------------------------------------------
    doc.add_heading("5. Configuración de la base de datos (.env)", 1)
    doc.add_paragraph(
        "La conexión a la base de datos SeisComP se define en el archivo .env de la "
        "instalación (se crea automáticamente). Valores por defecto:"
    )
    codigo(
        doc,
        "NEWPT_DB_HOST=10.54.217.69\n"
        "NEWPT_DB_DATABASE=seiscomp\n"
        "NEWPT_DB_USER=sysop\n"
        "NEWPT_DB_PASSWORD=sysop",
    )
    doc.add_paragraph("Edítalo si cambian los datos de conexión (tiene permisos 600).")

    # ------------------------------------------------------------------
    doc.add_heading("6. Cómo agregar nuevos perfiles locales (opción recomendada)", 1)
    doc.add_paragraph(
        "Los perfiles se detectan dinámicamente en cada evento: el graficador escanea la "
        "carpeta grillas/ buscando archivos slabP<número>.tmp y elige el más cercano al "
        "evento (si está a menos de 1.5°). Por eso, agregar perfiles NO requiere reinstalar "
        "ni recompilar: basta copiar los archivos."
    )
    doc.add_paragraph("Para agregar un nuevo perfil (por ejemplo, hacia el sur):")
    for t in [
        "Copiar los dos archivos a la carpeta de datos de la instalación "
        "(~/newpt/grillas/) y también a la carpeta grillas/ de la máquina de desarrollo:",
        "slabP0NN.tmp  y  topoP0NN.tmp  (el MISMO número en ambos; estilo de 3 dígitos, "
        "sin chocar con los existentes; hoy hay hasta P032, los nuevos serían P033 en adelante).",
        "Mantener el mismo formato: líneas # de comentario + columnas separadas por espacios "
        "(longitud col. 2, latitud col. 3, profundidad col. 6).",
        "El siguiente evento ya usará el perfil nuevo automáticamente; no hay que tocar nada más.",
    ]:
        doc.add_paragraph(t, style="List Bullet")
    codigo(
        doc,
        "scp slabP033.tmp topoP033.tmp usuario@maquina:~/newpt/grillas/",
    )
    doc.add_paragraph(
        "Los archivos de grillas/originales/ no interfieren: el escaneo solo mira el nivel "
        "raíz de grillas/. Si luego se genera el paquete con crear_instalador.sh, las "
        "grillas/ actualizadas quedan incluidas para instalaciones nuevas."
    )

    # ------------------------------------------------------------------
    doc.add_heading("7. Solución de problemas y notas", 1)
    for t in [
        "pip necesita internet o un mirror: si la máquina no tiene acceso, solicitar a TI "
        "un mirror de PyPI o un paquete offline de wheels.",
        "Si pip install cartopy no encuentra wheel para el Python instalado: solicitar a TI "
        "python3-cartopy (sudo apt) y crear el venv con --system-site-packages.",
        "PyInstaller requiere que el Python tenga la librería compartida (libpython3.X.so). "
        "El instalador lo detecta automáticamente: si python3 no la tiene (p. ej. Python "
        "compilado a mano sin --enable-shared, o de pyenv/conda), usa un intérprete del "
        "sistema que sí sirva (por ejemplo /usr/bin/python3.12) sin intervención. También "
        "se puede forzar con:  PYTHON_INTERPRETER=/usr/bin/python3 ./instalar_newpt.sh ...",
        "Reinstalar sobre una carpeta que ya tenía NewPT es seguro: se reconstruyen .dev/ y bin/.",
        "Si al abrir la ventana gráfica aparece 'No module named PIL._tkinter_finder': es un "
        "módulo de Pillow que el hook de PyInstaller excluye (excluye tkinter). El instalador "
        "ya lo fuerza en la compilación con --hidden-import tkinter y "
        "--hidden-import PIL._tkinter_finder.",
        "Todas las dependencias de Python van al venv con versiones EXACTAMENTE FIJADAS a las "
        "probadas en desarrollo: numpy 2.4.4, matplotlib 3.10.9, cartopy 0.25.0, pyproj 3.7.2, "
        "Pillow 12.2.0, psycopg2-binary 2.9.12, adjustText 1.4.0 y pyinstaller 6.22.2. Al "
        "terminar, el instalador imprime las versiones realmente instaladas.",
        "Si alguna vez aparece 'shapely.errors.GEOSException: Points of LinearRing do not form "
        "a closed linestring': es una incompatibilidad de matplotlib 3.11 con la graticula "
        "(gridlines) de cartopy. Por eso todas las versiones van fijadas (matplotlib 3.10.9); "
        "no actualizarlas por separado.",
        "El instalador nunca ejecuta apt y nunca instala nada en el Python del sistema.",
        "Distribuciones RHEL/CentOS/Rocky: los equivalentes son "
        "sudo dnf install -y python3 python3-virtualenv tkinter xterm.",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    # ------------------------------------------------------------------
    doc.add_heading("8. Anexo: lista para el encargado de TI", 1)
    doc.add_paragraph("El archivo LISTA_PARA_TI.txt que viaja con el instalador contiene:")
    codigo(
        doc,
        "sudo apt update\nsudo apt install -y python3 python3-venv python3-tk xterm",
    )
    doc.add_paragraph(
        "Entregar esa lista a TI junto con el paquete newpt_instalador.tar.gz. "
        "TI instala los 4 programas una sola vez; la instalación de NewPT la hace el "
        "usuario sin sudo."
    )

    doc.save(SALIDA)
    print("Generado:", SALIDA)


if __name__ == "__main__":
    main()