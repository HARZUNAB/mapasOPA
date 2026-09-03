#!/usr/bin/env python3
"""Genera INFORME_TESTS_MARCHA_BLANCA.docx con el detalle completo de las
pruebas unitarias de la marcha blanca de NewPT."""

import os
import sys
from datetime import date

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
except ImportError:
    print("Falta el módulo 'docx'. Instálalo y vuelve a ejecutar:")
    print("    pip install python-docx")
    sys.exit(1)

DIR = os.path.dirname(os.path.abspath(__file__))
SALIDA = os.path.join(DIR, "INFORME_TESTS_MARCHA_BLANCA.docx")


def codigo(doc, texto):
    """Bloque de comando en fuente monoespaciada."""
    for linea in texto.rstrip("\n").split("\n"):
        p = doc.add_paragraph()
        r = p.add_run(linea)
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)


def nota(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run("NOTA: " + texto)
    r.bold = True
    r.font.size = Pt(10)
    return p


def tabla(doc, encabezados, filas):
    t = doc.add_table(rows=1, cols=len(encabezados))
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    hdr = t.rows[0].cells
    for i, h in enumerate(encabezados):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8.5)
    for fila in filas:
        celdas = t.add_row().cells
        for i, v in enumerate(fila):
            celdas[i].text = str(v)
            for p in celdas[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8.5)
    doc.add_paragraph()
    return t


def main():
    doc = Document()
    doc.add_heading("Informe de pruebas — Marcha blanca NewPT", 0)
    doc.add_paragraph(
        "Detalle completo de las pruebas unitarias que se ejecutan antes de instalar "
        "el proyecto NewPT (contexto sismotectónico de eventos sísmicos), como control "
        "de calidad previo a la compilación de binarios."
    )
    doc.add_paragraph("Generado el " + date.today().strftime("%d/%m/%Y") + ".")

    # ------------------------------------------------------------------
    doc.add_heading("1. Objetivo y alcance", 1)
    doc.add_paragraph(
        "La marcha blanca valida, SIN conexión a la base de datos y SIN red, que la "
        "lógica crítica de NewPT funciona correctamente antes de instalar o compilar "
        "nada. Cubre tres suites:"
    )
    for t in [
        "Suite A — capturar.py: cálculos geográficos, parseo de la línea de evento "
        "y el manejo del perfil 3D (botón, zoom con la rueda y cierre coordinado).",
        "Suite B — consulta_evento.py: formato de salida hacia "
        "evento_data.txt, usando una base de datos SIMULADA (mock).",
        "Suite C — archivos de datos: consistencia de grillas, catálogo histórico, "
        "localidades y relieve.",
        "Además, valida en conjunto las versiones FIJADAS de los módulos de Python "
        "(numpy, matplotlib, cartopy, etc.), que es donde aparecieron problemas de "
        "compatibilidad en el pasado.",
    ]:
        doc.add_paragraph(t, style="List Bullet")
    nota(doc,
         "Los tests usan únicamente la librería estándar unittest: no agregan ninguna "
         "dependencia nueva al proyecto.")

    # ------------------------------------------------------------------
    doc.add_heading("2. Entorno de ejecución", 1)
    for t in [
        "Intérprete: el venv de la instalación (o uno temporal si solo se prueban).",
        "MPLBACKEND=Agg: matplotlib sin ventanas gráficas (entorno sin pantalla).",
        "NEWPT_DATA_DIR: apunta a la carpeta de datos a validar.",
        "PYTHONPATH: apunta a fuentes/ para poder importar los módulos.",
        "No se conecta a PostgreSQL: la conexión se simula con unittest.mock.",
        "No requiere acceso a internet una vez instalados los módulos fijados.",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    # ------------------------------------------------------------------
    doc.add_heading("3. Suite A — capturar.py (test_capturar.py)", 1)
    doc.add_paragraph(
        "Prueba la lógica pura del graficador: geodesia, parseo robusto y el manejo "
        "del perfil 3D (botón 'Ver Perfil 3D', zoom con la rueda del mouse y cierre "
        "coordinado de la ventana 2D/3D). En total la suite completa suma 31 tests "
        "(20 de capturar.py, 5 de consulta_evento.py y 6 de datos)."
    )
    tabla(
        doc,
        ["ID", "Objetivo", "Entrada", "Resultado esperado"],
        [
            ["A1", "Distancia nula", "Un punto comparado consigo mismo", "0.0 km"],
            ["A2", "Distancia real conocida",
             "Santiago (-70.65, -33.45) vs Valparaíso (-71.62, -33.05)",
             "≈103 km (tolerancia ±15 km)"],
            ["A3", "Simetría de la distancia", "d(A,B) y d(B,A)", "Ambos iguales"],
            ["A4", "Rumbo al norte", "Destino más al norte", "'N'"],
            ["A5", "Rumbo al sur", "Destino más al sur", "'S'"],
            ["A6", "Rumbo al este", "Destino más al este", "'E'"],
            ["A7", "Rumbo al oeste", "Destino más al oeste", "'W'"],
            ["A8", "Resolución de datos por entorno", "Variable NEWPT_DATA_DIR definida",
             "ruta_datos() devuelve esa carpeta"],
            ["A9", "Parseo de línea válida",
             "Línea de 19 campos con '35.42 S', '71.62 W', '95.3 km'",
             "dict con lat=-35.42, lon=-71.62, prof=95.3, fecha, mag y event_id correctos"],
            ["A10", "Magnitud sin dato",
             "Campo de magnitud vacío en la línea",
             "mag=0.0 y texto_magnitud 'M s/d'"],
            ["A11", "Profundidad corrupta",
             "Campo de profundidad no numérico",
             "prof=10.0 (valor por defecto) y aviso en pantalla"],
            ["A12", "Línea incompleta", "Menos de 12 campos separados por ';'",
             "Línea rechazada (None)"],
            ["A13", "Línea ajena al proyecto",
             "Texto sin la marca 'csn_'", "Línea rechazada"],
            ["A14", "Hemisferio norte/este", "'35.42 N' y '71.62 E'",
             "lat=+35.42 y lon=+71.62 (positivos)"],
            ["A15", "El botón abre el perfil 3D", "Ventana 2D plotear_evento renderizada, "
             "clic en el botón (fig._abrir_perfil_3d)",
             "Se abre una ventana 3D nueva: _VENTANA_3D_ABIERTA = True y hay 2D + 3D abiertas"],
            ["A16", "No abre una segunda 3D", "3D ya abierta y nuevo clic en el botón",
             "Mensaje 'Ya hay una ventana 3D abierta' y el número de figuras no cambia"],
            ["A17", "Reabre tras cerrar la 3D", "Se cierra la 3D y se hace clic de nuevo",
             "Se abre una 3D nueva (una sola a la vez, reabrible)"],
            ["A18", "Zoom con la rueda del mouse", "fig3d._zoom_3d({button:'up'}) y "
             "luego {button:'down'}", "Rango del eje X se reduce al acercar y aumenta al alejar"],
            ["A19", "Cerrar el 2D cierra el 3D", "Cerrar la ventana 2D principal "
             "(fig2d._al_cerrar_2d)", "La 3D se cierra con la 2D y _FIGURA_3D_ACTIVA = None"],
            ["A20", "Guarda autocorregible sin close_event",
             "Cerrar la 3D sin evento de cierre (backend Agg) y reclicar",
             "Se detecta que la figura 3D ya no existe y se abre una 3D nueva "
             "(sin banderas desincronizadas)"],
        ],
    )

    # ------------------------------------------------------------------
    doc.add_heading("4. Suite B — consultas SeisComP (test_consulta.py)", 1)
    doc.add_paragraph(
        "Prueba el formato EXACTO que consulta_evento.py escribe en evento_data.txt, "
        "que es el contrato con capturar.py. La conexión a PostgreSQL se simula "
        "(mock): no hace falta servidor ni red."
    )
    tabla(
        doc,
        ["ID", "Objetivo", "Entrada", "Resultado esperado"],
        [
            ["B1", "Estructura de la línea", "Fila simulada de la BD",
             "evento_data.txt con exactamente 19 campos separados por ';'"],
            ["B2", "Formato de fecha", "Origen con m_time_value dado",
             "OT escrita como '%Y-%m-%d %H:%M:%S'"],
            ["B3", "Coordenadas con cardinal", "lat=-33.45, lon=-71.62",
             "'33.45 S' y '71.62 W' (valor absoluto + sufijo)"],
            ["B4", "Profundidad formateada", "profundidad_km=95.3",
             "'95 km' (se trunca a entero)"],
            ["B5", "ID inexistente", "cursor sin filas",
             "NO se escribe archivo y la función termina sin error"],
            ["B6", "Limpieza preventiva", "evento_data.txt viejo preexistente",
             "Se elimina ANTES de consultar (evita ploteos fantasma)"],
            ["B7", "Magnitud no calculada", "magnitud=None en la fila",
             "Campo de magnitud vacío en la línea (aguas abajo será 'M s/d')"],
            ["B8", "Origen no preferido",
             "Origen sin evento asociado (no confirmado)",
             "Mensaje de ATENCIÓN al operador y NO se escribe archivo"],
        ],
    )

    # ------------------------------------------------------------------
    doc.add_heading("5. Suite C — archivos de datos (test_datos.py)", 1)
    doc.add_paragraph(
        "Valida la integridad de los datos que viajan con el instalador, leyendo "
        "desde NEWPT_DATA_DIR (en desarrollo: la carpeta ploteo/; en instalación: datos/)."
    )
    tabla(
        doc,
        ["ID", "Objetivo", "Verificación", "Resultado esperado"],
        [
            ["C1", "Pares de perfiles consistentes",
             "Por cada slabP###.tmp debe existir topoP###.tmp con el MISMO número",
             "Al menos un par; ningún huérfano crítico sin su compañero"],
            ["C2", "Formato de perfiles",
             "Cabecera '#' + filas numéricas con ≥7 columnas",
             "Todos los archivos parsean sin excepciones ('nan' tolerado en col. 7)"],
            ["C3", "Catálogo histórico",
             "base_2023_2026.dat: columnas 2 a 5 numéricas",
             "≥1 línea válida y todas parsean como flotantes"],
            ["C4", "Localidades", "localidades.csv con encabezado",
             "Existe y tiene ≥1 fila de localidad"],
            ["C5", "Relieve de fondo", "NE2_LR_LC_SR_W_DR.tif",
             "Existe y pesa más de 0 bytes"],
            ["C6", "Grillas globales binarias",
             "slab2_global.npy y topo_chile.npy (opcionales)",
             "Si existen, cargan con numpy sin errores y tienen 3 columnas"],
        ],
    )

    # ------------------------------------------------------------------
    doc.add_heading("6. Formas de ejecución", 1)
    doc.add_paragraph("a) Marcha blanca pura, sin instalar nada (crea un venv temporal, "
                      "instala las versiones fijadas, corre los 3 suites y limpia):")
    codigo(doc, "./instalar_newpt.sh --probar")
    doc.add_paragraph("b) Control automático dentro de la instalación: después de crear el "
                      "venv y verificar las versiones instaladas, y ANTES de compilar los "
                      "binarios, el instalador ejecuta los mismos tests. Si alguno falla:")
    codigo(doc, "============================================================\n"
                "  INSTALACIÓN DETENIDA: la marcha blanca (tests) falló.\n"
                "============================================================")
    doc.add_paragraph("y la instalación NO continúa hasta corregir el problema.")
    doc.add_paragraph("c) Ejecución manual de un suite específico (dentro del venv):")
    codigo(doc, "PYTHONPATH=fuentes MPLBACKEND=Agg \\\n"
                "python -m pytest fuentes/pruebas -v\n"
                "# solo un archivo:\n"
                "python -m pytest fuentes/pruebas/test_capturar.py -v")
    nota(doc, "La marcha blanca del instalador ejecuta los suites con pytest, "
              "que corre tanto tests unittest (como estos) como tests en estilo "
              "pytest puro que se agreguen a fuentes/pruebas/ del paquete de "
              "instalación (generado por crear_instalador.sh desde "
              "pruebas_uni/).")

    # ------------------------------------------------------------------
    doc.add_heading("7. Interpretación de resultados", 1)
    for t in [
        "OK: todos los tests pasan → la instalación sigue (compilación de binarios).",
        "FAIL: unittest lista cada caso fallido con nombre exacto (ej. "
        "test_a11_profundidad_corrupta) y el motivo.",
        "Ante un FAIL: revisar el caso listado, corregir el código o los datos, y volver "
        "a ejecutar ./instalar_newpt.sh (o --probar para iterar rápido).",
        "Errores de IMPORT en los tests suelen indicar un problema con las versiones de "
        "los módulos fijados: no relajar las versiones sin probar el conjunto completo.",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    doc.save(SALIDA)
    print("Generado:", SALIDA)


if __name__ == "__main__":
    main()
