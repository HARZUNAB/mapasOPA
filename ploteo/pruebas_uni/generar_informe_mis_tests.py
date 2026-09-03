#!/usr/bin/env python3
"""Genera INFORME_DETALLADO_MIS_TESTS_M1_M5.docx con el detalle completo de las
pruebas unitarias escritas de forma guiada (M1-M5) para el proyecto NewPT.

Los tests viven en pruebas_uni/mis_pruebas/ y están documentados aquí de forma
separada de los tests base del proyecto (suite IA).
"""

import os
import sys
from datetime import date

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:
    print("Falta el módulo 'docx'. Instálalo y vuelve a ejecutar:")
    print("    pip install python-docx")
    sys.exit(1)

DIR = os.path.dirname(os.path.abspath(__file__))
SALIDA = os.path.join(DIR, "mis_pruebas", "INFORME_DETALLADO_MIS_TESTS_M1_M5.docx")


def nota(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run("NOTA: " + texto)
    r.bold = True
    r.font.size = Pt(10)
    return p


def tabla(doc, encabezados, filas):
    t = doc.add_table(rows=1, cols=len(encabezados))
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    hdr = t.rows[0].cells
    for i, h in enumerate(encabezados):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8.5)
    for fila in filas:
        celdas = t.add_row().cells
        for i, v in enumerate(fila):
            celdas[i].text = str(v)
            for p in celdas[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8.5)
    doc.add_paragraph()
    return t


# =========================================================================
# Misiones (cada una es una función objetivo del proyecto)
# =========================================================================
MISIONES = [
    ("M1", "consultar_por_fecha_creacion (consulta_evento.py)",
     "test_m1_consulta.py",
     "Conexión simulada (mock) a la base de datos SeisComP. Verifica la consulta "
     "por origen, el fallback por evento, la limpieza preventiva y el formato "
     "exacto de los 19 campos que se escriben en evento_data.txt.",
     [
        ["M1_1", "Origen preferido escribe archivo",
         "Origen preferido: se escribe evento_data.txt con exactamente 19 campos "
         "separados por ';', con fecha, magnitud, coordenadas con hemisferio S/W, "
         "profundidad truncada a entero y event_id al final.", "PASSED"],
        ["M1_2", "Origen no preferido avisa y no escribe",
         "Origen NO preferido (id_evento=None): la función imprime el aviso "
         "'NO es la solución preferida' y NO genera evento_data.txt.", "PASSED"],
        ["M1_3", "ID inexistente no escribe archivo",
         "ID inexistente (fetchone=None en ambas consultas): no se escribe archivo "
         "y se imprime 'No se encontró'.", "PASSED"],
        ["M1_4", "Limpieza preventiva borra archivo viejo",
         "Si existía un evento_data.txt de una corrida anterior, se borra ANTES de "
         "conectar a la BD, evitando ploteos fantasmas.", "PASSED"],
        ["M1_5", "Magnitud None deja campo vacío",
         "Magnitud=None en la fila: el campo magnitud queda vacío ('') y el tipo "
         "queda 'M' por defecto (aguas abajo capturar.py lo interpreta como 'M s/d').", "PASSED"],
        ["M1_6", "Fallback por evento",
         "Si el ID no se encuentra como origen, se busca como evento usando la "
         "segunda consulta SQL; si se encuentra, se escribe el archivo correctamente.", "PASSED"],
        ["M1_7", "Región ausente usa 'Unknown Region'",
         "eventdescription ausente (region=None): el campo region se escribe como "
         "'Unknown Region'.", "PASSED"],
        ["M1_8", "Profundidad cero",
         "Profundidad 0.0 km: se escribe '0 km' (sin decimales), sin errores de "
         "formateo.", "PASSED"],
     ]),

    ("M2", "parsear_linea_evento (capturar.py)",
     "test_m2_capturar.py",
     "Conversión de la línea de texto de evento_data.txt a un dict. Verifica la "
     "extracción de lat/lon con hemisferios, el manejo robusto de magnitud y "
     "profundidad, y la validación de líneas inválidas.",
     [
        ["M2_1", "Línea válida campos correctos",
         "Línea válida completa de 19 campos: devuelve dict con fecha, "
         "lat=-35.42 (S→negativo), lon=-71.62 (W→negativo), prof=95.3, mag=4.6, "
         "tipo_mag=Ml, texto_magnitud='4.6 Ml' y event_id correctos.", "PASSED"],
        ["M2_2", "Magnitud vacía da 'M s/d'",
         "Magnitud vacía en la línea: campo mag=0.0 y texto_magnitud='M s/d (Ml)'.", "PASSED"],
        ["M2_3", "Profundidad corrupta usa default",
         "Profundidad corrupta ('Moments' en vez de número): valor por defecto "
         "prof=10.0 km e imprime aviso [Aviso].", "PASSED"],
        ["M2_4", "Línea incompleta devuelve None",
         "Línea incompleta (< 12 campos separados por ';'): la función la rechaza "
         "y devuelve None.", "PASSED"],
        ["M2_5", "Sin marca 'csn_' devuelve None",
         "Sin marca 'csn_' en el event_id (texto ajeno, vacío o None): la función "
         "la rechaza y devuelve None.", "PASSED"],
        ["M2_6", "Hemisferio norte latitud positiva",
         "Hemisferio Norte ('35.42 N'): la latitud se conserva positiva (+35.42).", "PASSED"],
        ["M2_7", "Hemisferio este longitud positiva",
         "Hemisferio Este ('71.62 E'): la longitud se conserva positiva (+71.62).", "PASSED"],
        ["M2_8", "Longitud 'O' es negativa",
         "Longitud con 'O' en vez de 'W' (español: Oeste): se interpreta como "
         "negativa, igual que 'W'.", "PASSED"],
        ["M2_9", "Profundidad 0.0 km",
         "Profundidad 0.0 km: se parsea como prof=0.0 sin errores.", "PASSED"],
     ]),

    ("M3", "lee_catalogo (lee_catalogo.py)",
     "test_m3_catalogo.py",
     "Lectura estandarizada del catálogo histórico (.dat). Verifica la corrección "
     "de columnas (lat en col 2, lon en col 3), el mapeo S/N a booleano, la opción "
     "sensibles=True y el manejo de archivos mínimos o vacíos.",
     [
        ["M3_1", "Latitud/longitud en columnas correctas",
         "Archivo con 2 eventos: lat (col 2) y lon (col 3) están en las columnas "
         "correctas (no invertidas); prof/mag/tipo se extraen bien.", "PASSED"],
        ["M3_2", "Múltiples eventos",
         "Múltiples eventos (3 filas): el DataFrame tiene exactamente 3 filas con "
         "todas parseadas correctamente.", "PASSED"],
        ["M3_3", "Fecha concatenada",
         "Fecha concatenada: df['fecha'] contiene col0 + ' ' + col1 en formato "
         "'YYYY-MM-DD HH:MM:SS'.", "PASSED"],
        ["M3_4", "Mapeo S/N a booleano",
         "Mapeo sensible: la columna 'S' se convierte a True y 'N' a False.", "PASSED"],
        ["M3_5", "sensibles=True numérico",
         "sensibles=True: la columna sensible se reemplaza por valores numéricos "
         "de la última columna del archivo.", "PASSED"],
        ["M3_6", "Archivo vacío",
         "Archivo vacío: numpy retorna arrays vacíos (con warnings) y la función "
         "devuelve un DataFrame de 0 filas sin excepción.", "PASSED"],
        ["M3_7", "Archivos independientes",
         "Archivos independientes: dos archivos diferentes producen DataFrames "
         "separados con datos no mezclados.", "PASSED"],
     ]),

    ("M4", "contar_lineas y procesar_xyz (preprocesa_grillas.py)",
     "test_m4_grillas.py",
     "Preprocesamiento de grillas topográficas/de slab (.xyz). Verifica el conteo "
     "de líneas, la conversión a array numpy, la ordenación por latitud, el filtro "
     "de filas y el manejo de comentarios, vacíos y separadores.",
     [
        ["M4_1", "contar_lineas: 3 líneas",
         "contar_lineas con 3 líneas escritas: retorna 3.", "PASSED"],
        ["M4_2", "contar_lineas: archivo vacío",
         "contar_lineas con archivo vacío: retorna 0 sin errores.", "PASSED"],
        ["M4_3", "contar_lineas: archivo inexistente",
         "contar_lineas con archivo inexistente: lanza FileNotFoundError u OSError.", "PASSED"],
        ["M4_4", "procesar_xyz: orden por latitud",
         "3 puntos con coordenadas desordenadas: el array resultante queda ordenado "
         "por latitud (columna 1) en orden ascendente.", "PASSED"],
        ["M4_5", "procesar_xyz: filtro descarta filas",
         "Filtro aplicado (prof > 150): solo se conservan las filas que pasan el "
         "filtro; las demás se descartan.", "PASSED"],
        ["M4_6", "procesar_xyz: filtro None conserva todo",
         "Filtro=None: todas las filas se conservan sin descartar.", "PASSED"],
        ["M4_7", "procesar_xyz: comentarios ignorados",
         "Líneas con comentarios (#): se ignoran; solo cuentan las líneas de datos.", "PASSED"],
        ["M4_8", "procesar_xyz: líneas vacías ignoradas",
         "Líneas vacías o con solo espacios: se ignoran correctamente.", "PASSED"],
        ["M4_9", "procesar_xyz: un solo dato",
         "1 solo dato: el array tiene exactamente 1 fila.", "PASSED"],
        ["M4_10", "procesar_xyz: separador whitespace",
         "Separador whitespace (espacios y tabs mezclados): sep=None maneja ambos "
         "correctamente.", "PASSED"],
     ]),

    ("M5", "_promediar_tris (capturar.py)",
     "test_m5_tris.py",
     "Agrupación de puntos (lon, lat, altitud/profundidad) en celdas de un tamaño "
     "dado (res) y cálculo del valor promedio por celda, para generar superficies "
     "3D suaves y ligeras.",
     [
        ["M5_1", "Promedio correcto por celda",
         "2 puntos en la misma celda (res=1.0): 1 celda resultante con el promedio "
         "correcto de z.", "PASSED"],
        ["M5_2", "Todos en una celda",
         "4 puntos en la misma celda: 1 celda con promedio "
         "(100+200+300+400)/4 = 250.0.", "PASSED"],
        ["M5_3", "res grande una sola celda",
         "res=10.0 (muy grande): puntos cercanos se agrupan en 1 sola celda.", "PASSED"],
        ["M5_4", "res pequeño cada punto es celda",
         "res=0.001 (muy pequeño): cada punto es su propia celda, sin promediar; "
         "los valores originales quedan intactos.", "PASSED"],
        ["M5_5", "Profundidades negativas",
         "Profundidades negativas (slab subductado): promedio negativo correcto "
         "(-100 y -300 → -200).", "PASSED"],
        ["M5_6", "Array vacío",
         "Arrays vacíos: salida vacía (3 arrays de longitud 0) sin errores.", "PASSED"],
        ["M5_7", "Un solo punto",
         "1 solo punto: 1 celda con el valor original sin promediar.", "PASSED"],
     ]),
]


def main():
    doc = Document()
    doc.add_heading("Informe detallado de mis pruebas unitarias (M1-M5)", 0)
    doc.add_paragraph(
        "Documento exclusivo de las pruebas unitarias desarrolladas de forma guiada "
        "para el proyecto NewPT. Estas 41 pruebas se mantienen SEPARADAS de los "
        "tests base del proyecto (suite sugerida por IA, documentada en su propio "
        "informe)."
    )
    doc.add_paragraph(
        "Los archivos viven en pruebas_uni/mis_pruebas/ con prefijo test_m*.py y "
        "comentarios explicativos paso a paso para facilitar la comprensión de "
        "cada caso."
    )
    doc.add_paragraph("Generado el " + date.today().strftime("%d/%m/%Y") + ".")

    total = 0
    for codigo, funcion, archivo, descripcion, filas in MISIONES:
        total += len(filas)

    doc.add_heading("Resumen", 1)
    nota(doc,
         f"En total este informe documenta {total} pruebas unitarias distribuidas "
         f"en {len(MISIONES)} misiones, todas con resultado PASSED.")

    # ------------------------------------------------------------------
    for codigo, funcion, archivo, descripcion, filas in MISIONES:
        doc.add_heading(f"{codigo} — {funcion} ({len(filas)} pruebas)", 1)
        doc.add_paragraph("Archivo: " + archivo)
        doc.add_paragraph(descripcion)
        tabla(doc, ["ID", "Caso que cubre", "Detalle", "Resultado"], filas)

    doc.add_heading("Conclusión", 1)
    doc.add_paragraph(
        "Las 5 misiones equivalen a la cobertura guiada de la lógica crítica de "
        "NewPT: consulta a BD (M1), parseo de la línea de evento (M2), lectura del "
        "catálogo (M3), preprocesamiento de grillas (M4) y generación de superficies "
        "3D (M5). Todas corren correctamente con:"
    )
    for t in [
        "python3 -m pytest pruebas_uni/mis_pruebas -v",
        "python3 -m pytest pruebas_uni -v   (suite completa: 72 pruebas)",
    ]:
        doc.add_paragraph(t, style="List Bullet")

    doc.save(SALIDA)
    print("Generado:", SALIDA)


if __name__ == "__main__":
    main()
