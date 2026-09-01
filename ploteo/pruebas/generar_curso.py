#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Genera CURSO_TESTS_UNITARIOS.docx: curso práctico de pytest + black para NewPT.

Las soluciones del laboratorio viven en SOLUCIONES y se usan para dos cosas:
  1) Renderizar los bloques de código del anexo.
  2) `python3 pruebas/generar_curso.py --verificar` las ejecuta de verdad con
     pytest (en /tmp) para garantizar que lo publicado funciona.

Requiere python-docx (venv vtest) para el docx, y pytest para --verificar.
"""
import os
import subprocess
import sys
import tempfile

DIR = os.path.dirname(os.path.abspath(__file__))
BUILD = "/tmp/opencode/curso_build"
MEDIA = os.path.join(DIR, "CURSO_MEDIA")
SALIDA = os.path.join(DIR, "CURSO_TESTS_UNITARIOS.docx")


# =========================================================================
# Soluciones oficiales del laboratorio (ejecutables de verdad)
# =========================================================================
SOLUCIONES = {
    "test_lab1_avance.py": '''import capturar


def test_porcentaje_y_bloques(capsys):
    """Al 50% la barra dibuja 10 bloques llenos y 10 vacíos."""
    capturar.mostrar_avance(50, 100, "leyendo grillas")
    fuera = capsys.readouterr().out
    assert "50%" in fuera
    assert "█" * 10 in fuera
    assert "-" * 10 in fuera


def test_inicio_sin_bloques(capsys):
    """Con paso=0 la barra arranca completamente vacía."""
    capturar.mostrar_avance(0, 100, "inicio")
    fuera = capsys.readouterr().out
    assert "0%" in fuera
    assert "█" not in fuera


def test_mensaje_incluido(capsys):
    """El mensaje del operador aparece junto a la barra."""
    capturar.mostrar_avance(25, 100, "cargando catálogo")
    assert "cargando catálogo" in capsys.readouterr().out
''',
    "test_lab2_contar_lineas.py": '''import pytest

import preprocesa_grillas


def test_cuenta_todas_las_lineas(tmp_path):
    """Cuenta encabezado y filas de datos: 3 líneas escritas, 3 contadas."""
    f = tmp_path / "mini.xyz"
    f.write_text("lon lat prof\\n-71.0 -35.0 50.0\\n-70.0 -36.0 60.0\\n")
    assert preprocesa_grillas.contar_lineas(str(f)) == 3


def test_archivo_vacio_da_cero(tmp_path):
    """Un archivo sin contenido cuenta 0 líneas, sin explotar."""
    f = tmp_path / "vacio.xyz"
    f.write_text("")
    assert preprocesa_grillas.contar_lineas(str(f)) == 0


def test_archivo_inexistente_lanza_oserror(tmp_path):
    """Contar en un archivo que no existe debe fallar de forma visible."""
    with pytest.raises(OSError):
        preprocesa_grillas.contar_lineas(str(tmp_path / "no_existe.txt"))
''',
    "test_lab3_inicializar_csv.py": '''import capturar


def test_csv_reiniciado_con_encabezado(tmp_path, monkeypatch, capsys):
    """inicializar_csv borra restos anteriores y escribe el encabezado exacto."""
    destino = tmp_path / "datos_seiscomp.csv"
    destino.write_text("basura de una corrida anterior")
    monkeypatch.setattr(capturar, "OUTPUT_FILE", str(destino))

    capturar.inicializar_csv()

    contenido = destino.read_text(encoding="utf-8")
    assert contenido.startswith(",Fecha_Hora,Latitud")
    assert "basura" not in contenido
    assert "CSV reiniciado" in capsys.readouterr().out
''',
    "test_lab4_lee_catalogo.py": '''import pytest

from lee_catalogo import lee_catalogo

CATALOGO = (
    "2026-01-02 03:04:05 -35.42 -71.62 95.0 4.6 Ml S 1\\n"
    "2026-02-03 04:05:06 -33.10 -70.20 10.0 3.2 Mw N 0\\n"
)


@pytest.fixture
def catalogo(tmp_path):
    """Catálogo mínimo de dos eventos en formato .dat."""
    f = tmp_path / "mini.dat"
    f.write_text(CATALOGO)
    return str(f)


def test_columnas_geograficas(catalogo):
    """Latitud es la columna 2 y longitud la 3 (formato .dat corregido)."""
    df = lee_catalogo(catalogo)
    assert list(df["lat"]) == [-35.42, -33.10]
    assert list(df["lon"]) == [-71.62, -70.20]


def test_fecha_unida(catalogo):
    """Fecha y hora se combinan en una sola columna 'fecha'."""
    df = lee_catalogo(catalogo)
    assert df["fecha"][0] == "2026-01-02 03:04:05"


def test_flag_sensible_por_letra(catalogo):
    """La columna S/N se mapea a True/False."""
    df = lee_catalogo(catalogo)
    assert list(df["sensible"]) == [True, False]


def test_sensibles_true_lee_columna_numerica(catalogo):
    """Con sensibles=True la columna final numérica reemplaza el flag."""
    df = lee_catalogo(catalogo, sensibles=True)
    assert list(df["sensible"]) == [1.0, 0.0]
''',
    "test_lab5_ruta_datos.py": '''import os

import consulta_evento


def test_prioridad_variable_entorno(monkeypatch, tmp_path):
    """NEWPT_DATA_DIR definida gana sobre cualquier otra ubicación."""
    monkeypatch.setenv("NEWPT_DATA_DIR", str(tmp_path))
    assert consulta_evento.ruta_datos() == str(tmp_path)


def test_sin_entorno_usa_carpeta_del_script(monkeypatch):
    """Sin variable de entorno, se usa el directorio del propio módulo."""
    monkeypatch.delenv("NEWPT_DATA_DIR", raising=False)
    esperado = os.path.dirname(os.path.abspath(consulta_evento.__file__))
    assert consulta_evento.ruta_datos() == esperado
''',
    "test_lab6_procesar_xyz.py": '''import pytest

import preprocesa_grillas


def _a_float(partes):
    return [float(p) for p in partes]


def test_orden_por_latitud(tmp_path):
    """procesar_xyz devuelve las filas ordenadas por latitud (col 1)."""
    f = tmp_path / "g.xyz"
    f.write_text("# comentario\\n-71.0 -35.0 50.0\\n-70.0 -36.0 60.0\\n")
    arr = preprocesa_grillas.procesar_xyz(str(f), " ", 3, _a_float, None, "demo")
    assert arr.shape == (2, 3)
    assert list(arr[:, 1]) == sorted(arr[:, 1])


def test_filtro_descarta_filas(tmp_path):
    """Solo pasan el filtro las filas cuya primera columna es menor que 5."""
    f = tmp_path / "g2.xyz"
    f.write_text("1 2 3\\n4 5 6\\n7 8 9\\n")
    arr = preprocesa_grillas.procesar_xyz(
        str(f), " ", 3, _a_float, lambda v: v[0] < 5, "demo"
    )
    assert arr.shape == (2, 3)


def test_error_de_lectura_se_propaga(tmp_path, monkeypatch):
    """Si contar_lineas falla, procesar_xyz no debe tragarse el error."""
    monkeypatch.setattr(
        preprocesa_grillas,
        "contar_lineas",
        lambda path: (_ for _ in ()).throw(OSError("disco lleno (simulado)")),
    )
    with pytest.raises(OSError, match="simulado"):
        preprocesa_grillas.procesar_xyz(
            "no_importa.xyz", " ", 3, _a_float, None, "demo"
        )
''',
}


# =========================================================================
# Utilidades docx
# =========================================================================
def codigo(doc, texto, size=9.5):
    for linea in texto.rstrip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(linea)
        r.font.name = "Consolas"
        r.font.size = Pt(size)


def imagen(doc, nombre, ancho=6.3, pie=None):
    path = os.path.join(BUILD, nombre)
    if not os.path.exists(path):
        print(f"[AVISO] falta imagen {path}")
        return
    doc.add_picture(path, width=Inches(ancho))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if pie:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(pie)
        r.italic = True
        r.font.size = Pt(9)


def video_ref(doc, archivo, descripcion):
    p = doc.add_paragraph()
    r = p.add_run("▶ VIDEO  ")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(f"CURSO_MEDIA/{archivo} — {descripcion}")
    r2.font.size = Pt(10)


def nota(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run("NOTA: " + texto)
    r.bold = True
    r.font.size = Pt(10)


def tabla(doc, encabezados, filas):
    t = doc.add_table(rows=1, cols=len(encabezados))
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    for i, h in enumerate(encabezados):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
    for fila in filas:
        cells = t.add_row().cells
        for i, v in enumerate(fila):
            cells[i].text = str(v)
    doc.add_paragraph()
    return t


# =========================================================================
# Contenido del curso
# =========================================================================
def construir(doc):
    # ---------------- portada ----------------
    t = doc.add_heading("CURSO DE TESTS UNITARIOS PARA NewPT", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("pytest + black · material práctico con capturas y videos")
    r.italic = True

    doc.add_paragraph(
        "Este curso te enseña a escribir tus propias pruebas unitarias para el "
        "proyecto NewPT usando pytest (framework de pruebas) y black (formato "
        "automático de código). Los 31 tests que ya existen quedan como red de "
        "seguridad: tú agregas los tuyos encima."
    )
    doc.add_heading("Cómo usar este material", level=2)
    for txt in (
        "Lee cada capítulo en orden; los ejemplos provienen del código real del "
        "proyecto (capturar.py, consulta_evento.py, lee_catalogo.py, "
        "preprocesa_grillas.py).",
        "Los videos están en la carpeta CURSO_MEDIA (junto a este documento). "
        "Ábrelos con doble clic cuando veas el marcador ▶ VIDEO.",
        "El capítulo 8 es el laboratorio: ahí escribes TUS tests. Las soluciones "
        "comentadas están en el capítulo 9 — intenta cada ejercicio antes de mirar.",
        "Regla de oro del proyecto: todo archivo test_*.py que dejes en pruebas/ "
        "corre automáticamente en la marcha blanca del instalador. Valida siempre "
        "en local antes (python3 -m pytest pruebas -v) para no bloquear "
        "instalaciones ajenas.",
    ):
        doc.add_paragraph(txt, style="List Bullet")

    doc.add_heading("Los videos del curso", level=2)
    tabla(
        doc,
        ["Archivo", "Qué muestra"],
        [
            ["v1_correr_suite.mp4", "Correr la suite completa y leer el resultado"],
            ["v2_primer_test.mp4", "Tu primer test, paso a paso desde cero"],
            ["v3_leer_traceback.mp4", "Leer un fallo (traceback) y corregirlo"],
            ["v4_black.mp4", "black formateando: antes y después"],
            ["v5_mock.mp4", "Un mock en acción (BD simulada)"],
        ],
    )

    # ---------------- cap 1 ----------------
    doc.add_page_break()
    doc.add_heading("1. Qué es un test unitario y por qué te sirve", level=1)
    doc.add_paragraph(
        "Un test unitario es un programita que llama a UNA función de tu proyecto "
        "con datos conocidos y verifica que la respuesta sea la esperada. Su "
        "trabajo principal es protegerte de las regresiones: errores que "
        "reaparecen en código que ya funcionaba."
    )
    doc.add_paragraph(
        "En NewPT esto no es teoría: los errores de latitud en la columna 8 y de "
        "profundidad truncada que ya sufriste son exactamente el tipo de bug que "
        "un test atrapa en segundos, antes de que llegue al mapa. Por eso el "
        "instalador ejecuta la suite como 'marcha blanca': si algo se rompe, la "
        "instalación se detiene antes de compilar binarios inútiles."
    )
    imagen(doc, "g_flujo.png", pie="El lugar de los tests dentro de instalar_newpt.sh")
    nota(doc, "un test unitario NO prueba la base de datos real, la red ni el "
              "mapa en pantalla: esas son pruebas de integración/manuales. Cada "
              "cosa a su nivel.")

    # ---------------- cap 2 ----------------
    doc.add_heading("2. Anatomía de un test", level=1)
    doc.add_paragraph(
        "Un test es una función normal cuyo nombre empieza con test_. pytest la "
        "descubre sola, la ejecuta y decide su destino según las aserciones que "
        "contenga. Nada más de magia."
    )
    imagen(doc, "g_anatomia.png", pie="Las partes de un test y qué hace cada una")
    doc.add_paragraph(
        "El proyecto ya tiene 31 tests escritos al estilo unittest (clases con "
        "assertEqual). pytest los ejecuta sin tocarlos. TUS tests nuevos irán en "
        "estilo pytest: más corto y directo. Compara:"
    )
    imagen(doc, "s_estilos.png", pie="Mismo test en ambos estilos — pytest es el de la derecha")
    tabla(
        doc,
        ["Convención", "Regla"],
        [
            ["Ubicación", "pruebas/ (junto a los tests existentes)"],
            ["Nombre de archivo", "test_loquesea.py (obligatorio test_ al inicio)"],
            ["Nombre de función", "def test_comportamiento(): (test_ al inicio)"],
            ["Agrupar (opcional)", "class TestAlgo: sin __init__"],
        ],
    )

    # ---------------- cap 3 ----------------
    doc.add_heading("3. Aserciones: el corazón de todo test", level=1)
    doc.add_paragraph(
        "En pytest no memorizas 30 métodos: usas el assert de Python y pytest te "
        "muestra, cuando falla, los valores reales de cada lado. Traducción de lo "
        "que ya conoces:"
    )
    tabla(
        doc,
        ["unittest (existente)", "pytest (tuyo)", "Cuándo usarlo"],
        [
            ["assertEqual(a, b)", "assert a == b", "igualdad exacta"],
            ["assertNotEqual(a, b)", "assert a != b", "diferencia"],
            ["assertTrue(x)", "assert x", "condición booleana"],
            ["assertIn(a, b)", "assert a in b", "subcadena o elemento"],
            ["assertIsNone(x)", "assert x is None", "ausencia de valor"],
            ["assertAlmostEqual(a, b)", "assert a == pytest.approx(b)", "floats"],
            ["assertRaises(E, f)", "with pytest.raises(E): f()", "excepciones esperadas"],
        ],
    )
    imagen(doc, "s_asserts.png", pie="Aserciones típicas sobre el código real de NewPT")

    # ---------------- cap 4 ----------------
    doc.add_heading("4. Fixtures: preparar y limpiar el escenario", level=1)
    doc.add_paragraph(
        "Muchos tests necesitan un escenario: archivos temporales, variables de "
        "entorno, datos de ejemplo. En unittest eso vive en setUp/tearDown (mira "
        "la clase BaseConsulta de pruebas/test_consulta.py). En pytest se usa "
        "un @pytest.fixture: una función que prepara el escenario y que cada test "
        "recibe como parámetro."
    )
    imagen(doc, "s_fixtures.png", pie="Tu BaseConsulta real y su equivalente pytest")
    tabla(
        doc,
        ["Fixture integrado de pytest", "Qué te da"],
        [
            ["tmp_path", "Carpeta temporal única para el test (objeto Path)"],
            ["capsys", "Captura lo que la función imprime en pantalla"],
            ["monkeypatch.setenv/delenv", "Cambia/restaura variables de entorno"],
            ["monkeypatch.setattr", "Reemplaza atributos o funciones y los restaura solo"],
        ],
    )
    nota(doc, "monkeypatch deshace cada cambio al terminar el test: tus pruebas "
              "nunca contaminan al vecino. Por eso no necesitas tearDown manual.")

    # ---------------- cap 5 ----------------
    doc.add_heading("5. Ejecutar los tests y leer los resultados", level=1)
    codigo(doc, "cd ploteo\npython3 -m pytest pruebas -v          # toda la carpeta\n"
                "python3 -m pytest pruebas/test_lab1_avance.py -v   # un archivo\n"
                "python3 -m pytest pruebas -k profundidad -v # por nombre\n"
                "python3 -m pytest pruebas -x               # detener en el 1er fallo\n"
                "python3 -m pytest pruebas --tb=short       # tracebacks cortos")
    imagen(doc, "s_suite.png", pie="Salida real: cada test en una línea + resumen final")
    doc.add_paragraph(
        "Leer un fallo: pytest agrupa los FAILURES y te muestra el assert exacto "
        "que falló con los valores de cada lado. FAIL significa que el código "
        "corrió pero dio algo inesperado; ERROR significa que ni siquiera pudo "
        "ejecutarse (excepción al preparar el escenario)."
    )
    imagen(doc, "s_fail.png", pie="Un FAIL bien leído se corrige solo")
    video_ref(doc, "v1_correr_suite.mp4", "correr toda la suite y entender cada línea")
    video_ref(doc, "v3_leer_traceback.mp4", "leer un traceback y corregir el test")

    # ---------------- cap 6 ----------------
    doc.add_heading("6. Mocks: simular lo lento o lo ajeno", level=1)
    doc.add_paragraph(
        "Un mock reemplaza una dependencia real (base de datos, red, reloj) por "
        "una versión falsa que controlas. Así pruebas TU lógica sin depender de "
        "servidores. La suite B ya lo hace con PostgreSQL: la conexión se "
        "sustituye por objetos falsos y aun así verifica el contrato completo de "
        "19 campos."
    )
    imagen(doc, "g_mock.png", pie="El test nunca toca la BD real")
    doc.add_paragraph(
        "La herramienta en pytest es monkeypatch.setattr: reemplaza un atributo "
        "durante el test y lo restaura al terminar. Esqueleto usado por tu suite B:"
    )
    codigo(doc,
           "def test_consulta_simulada(monkeypatch):\n"
           "    conexion_falsa = FakeConexion(FakeCursor([FILA_ORIGEN]))\n"
           "    monkeypatch.setattr('psycopg2.connect',\n"
           "                        lambda *a, **k: conexion_falsa)\n"
           "    consulta_evento.consultar_por_fecha_creacion('csn_sc62026nkkbb')\n"
           "    # ... y ahora se verifica el archivo escrito como si nada")
    video_ref(doc, "v5_mock.mp4", "el mock en acción, de la idea al PASSED")
    nota(doc, "usa mocks SOLO para dependencias externas (BD, red, tiempo, "
              "archivos gigantes). Mockear la función que estás probando no "
              "prueba nada.")

    # ---------------- cap 7 ----------------
    doc.add_heading("7. Buenas prácticas y el flujo con black", level=1)
    imagen(doc, "g_AAA.png", pie="Estructura mental de todo test: el ciclo AAA")
    for txt in (
        "Un concepto por test: si el nombre tiene una 'y' dudosa, son dos tests.",
        "Nombres que cuentan el comportamiento: test_origen_no_preferido_no_escribe_archivo.",
        "Primero hazlo fallar a propósito (cambia el esperado) para confirmar que "
        "el test realmente observa algo.",
        "No mockees lo que controlas; no pruebes detalles internos, prueba el "
        "resultado visible.",
        "plotear_evento() NO se testea unitariamente: dibuja mapas con cartopy. "
        "Eso se valida a mano con un evento real. Saber qué NO testear también "
        "es parte del oficio.",
    ):
        doc.add_paragraph(txt, style="List Bullet")
    doc.add_heading("El flujo de trabajo completo", level=2)
    codigo(doc,
           "1. escribe el test (rojo: aún no pasa o no existe el comportamiento)\n"
           "2. python3 -m pytest pruebas/test_tuyo.py -v   → PASSED\n"
           "3. black pruebas/test_tuyo.py                  → formato uniforme\n"
           "4. agrega el docstring que explique qué protege el test\n"
           "5. python3 -m pytest pruebas -v                → todo verde = listo")
    imagen(doc, "s_black.png", pie="black antes y después: tú escribes, black ordena")
    video_ref(doc, "v4_black.mp4", "black en acción")
    nota(doc, "black solo se aplica a TUS archivos de pruebas; el código de "
              "producción (capturar.py, etc.) no se reformatea para no generar "
              "diferencias enormes.")

    # ---------------- cap 8: laboratorio ----------------
    doc.add_page_break()
    doc.add_heading("8. Laboratorio: escribe los tests del proyecto", level=1)
    doc.add_paragraph(
        "Seis ejercicios progresivos que cubren TODO módulo aún sin tests. Crea "
        "cada archivo dentro de pruebas/ con el nombre indicado. Al terminar cada "
        "uno: pytest en verde → black → docstrings (ya incluidos en las "
        "soluciones del capítulo 9)."
    )
    tabla(
        doc,
        ["#", "Archivo a crear", "Función objetivo", "Aprendes"],
        [
            ["L1", "test_lab1_avance.py", "capturar.mostrar_avance", "capsys"],
            ["L2", "test_lab2_contar_lineas.py", "preprocesa_grillas.contar_lineas", "tmp_path, errores"],
            ["L3", "test_lab3_inicializar_csv.py", "capturar.inicializar_csv", "monkeypatch.setattr"],
            ["L4", "test_lab4_lee_catalogo.py", "lee_catalogo.lee_catalogo", "@pytest.fixture"],
            ["L5", "test_lab5_ruta_datos.py", "consulta_evento.ruta_datos", "monkeypatch.setenv"],
            ["L6", "test_lab6_procesar_xyz.py", "preprocesa_grillas.procesar_xyz", "mock de collaborator"],
        ],
    )
    video_ref(doc, "v2_primer_test.mp4", "hazlo primero con L1 siguiendo el video")

    labs = [
        ("L1 — mostrar_avance (capturar.py)",
         "pruebas/test_lab1_avance.py",
         "La barra de progreso escribe en stdout con sys.stdout.write. Con capsys "
         "capturas esa salida y verificas porcentaje, bloques █ y el mensaje.",
         ["50 de 100 → debe contener '50%', 10 '█' y 10 '-'",
          "0 de 100 → '0%' y ningún bloque lleno",
          "El mensaje del operador aparece en la salida"],
         "s_l1.png"),
        ("L2 — contar_lineas (preprocesa_grillas.py)",
         "pruebas/test_lab2_contar_lineas.py",
         "Función pura de archivos: cuenta líneas leídas en binario. Ideal para "
         "tmp_path: escribes archivos mínimos que pytest borra solo.",
         ["3 líneas escritas → 3 contadas (incluye el encabezado)",
          "Archivo vacío → 0, sin explotar",
          "Archivo inexistente → pytest.raises(OSError)"],
         "s_l2.png"),
        ("L3 — inicializar_csv (capturar.py)",
         "pruebas/test_lab3_inicializar_csv.py",
         "Reinicia el CSV de SeisComP. Ojo: OUTPUT_FILE se calcula al importar, "
         "por eso aquí se reemplaza con monkeypatch.setattr(capturar, "
         "'OUTPUT_FILE', ...) en lugar de cambiar el entorno.",
         ["Escribe basura previa → tras inicializar queda solo el encabezado",
          "El encabezado empieza con ',Fecha_Hora,Latitud'",
          "Imprime '[INFO] CSV reiniciado' (capsys)"],
         None),
        ("L4 — lee_catalogo (lee_catalogo.py)",
         "pruebas/test_lab4_lee_catalogo.py",
         "Estandariza la lectura del catálogo histórico (.dat). Necesitas un "
         "fixture con un catálogo mínimo: fecha, hora, lat(col2), lon(col3), "
         "prof(col4), mag(col5), tipo(col6), S/N(col7) y numérico final.",
         ["lat/lon caen en las columnas correctas",
          "fecha y hora se unen en df['fecha']",
          "S/N se mapea a True/False; con sensibles=True se usa la columna numérica final"],
         None),
        ("L5 — ruta_datos (consulta_evento.py)",
         "pruebas/test_lab5_ruta_datos.py",
         "La misma lógica de rutas que ya probamos en capturar, pero en el módulo "
         "de consulta. Practica monkeypatch.setenv/delenv.",
         ["Con NEWPT_DATA_DIR definida → devuelve exactamente ese valor",
          "Sin la variable → devuelve el directorio del propio módulo"],
         None),
        ("L6 — procesar_xyz (preprocesa_grillas.py)",
         "pruebas/test_lab6_procesar_xyz.py",
         "El preprocesador de grillas: convierte, filtra y ordena por latitud. "
         "Aquí además practicas mock: reemplazas contar_lineas por una versión "
         "que falla y verificas que el error se propaga.",
         ["Filas ordenadas por la columna 1 (latitud)",
          "El filtro descarta las filas que no pasan",
          "contar_lineas simulada con OSError → pytest.raises(OSError, match='simulado')"],
         None),
    ]
    for titulo, archivo, intro, criterios, img in labs:
        doc.add_heading(titulo, level=2)
        doc.add_paragraph(archivo, style="Intense Quote")
        doc.add_paragraph(intro)
        doc.add_paragraph("Criterios de aceptación:", style="List Bullet")
        for c in criterios:
            doc.add_paragraph(c, style="List Bullet 2")
        if img:
            imagen(doc, img)

    # ---------------- cap 9: soluciones ----------------
    doc.add_page_break()
    doc.add_heading("9. Anexo: soluciones comentadas", level=1)
    doc.add_paragraph(
        "Intenta cada ejercicio antes de mirar. Estas soluciones están formateadas "
        "con black y cada test lleva su docstring explicando qué protege — así "
        "deben quedar los tuyos. (Fueron ejecutadas y validadas con pytest antes "
        "de publicar este documento.)"
    )
    for nombre in sorted(SOLUCIONES):
        doc.add_heading(nombre, level=2)
        codigo(doc, SOLUCIONES[nombre], size=9)
    doc.add_paragraph()
    nota(doc, "cuando tus tests estén verdes, quedan integrados solos: pytest "
              "los descubre en pruebas/ y la marcha blanca del instalador los "
              "correrá de ahí en adelante.")


def verificar():
    """Ejecuta las soluciones con pytest en /tmp para garantizar que funcionan."""
    from shutil import which
    pytest_bin = which("pytest")
    if pytest_bin:
        base = [pytest_bin]
    else:  # consola pytest fuera de PATH: usar módulo
        base = [sys.executable, "-m", "pytest"]
    with tempfile.TemporaryDirectory(prefix="curso_lab_") as tmp:
        for nombre, codigo_src in SOLUCIONES.items():
            with open(os.path.join(tmp, nombre), "w", encoding="utf-8") as f:
                f.write(codigo_src)
        env = dict(os.environ)
        env["PYTHONPATH"] = os.path.dirname(DIR)
        env["MPLBACKEND"] = "Agg"
        r = subprocess.run(
            base + [tmp, "-v", "--tb=short"],
            env=env, text=True, capture_output=True)
        print(r.stdout[-3000:])
        if r.returncode != 0:
            sys.exit("[ERROR] las soluciones NO pasan: revisa el anexo")
        print("SOLUCIONES VERIFICADAS: todo PASSED")


if __name__ == "__main__":
    if "--verificar" in sys.argv:
        verificar()
    else:
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            print("Instala python-docx:  pip install python-docx")
            sys.exit(1)
        doc = Document()
        construir(doc)
        doc.save(SALIDA)
        print("Generado:", SALIDA)
